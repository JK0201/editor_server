import io
from datetime import datetime

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Twips
from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document
from app.models.script_line import ScriptLine
from app.models.speaker import Speaker


def _add_content(
    doc: DocxDocument,
    document: Document,
    speakers: dict,
    lines: list[ScriptLine],
):
    dt = datetime.strptime(
        f"{document.recorded_date}{document.recorded_time}", "%y%m%d%H%M%S"
    )
    formatted_datetime = dt.strftime("%y.%m.%d_%H:%M:%S")

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(formatted_datetime)
    date_run.bold = True
    date_run.underline = True

    for line in lines:
        speaker_name = speakers.get(line.speaker_id, "Unknown")
        para = doc.add_paragraph()

        # 탭 스탑 설정 (twips 단위, **직접 해보면서 조정 필요**)
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab1 = OxmlElement("w:tab")
        tab1.set(qn("w:val"), "left")
        tab1.set(qn("w:pos"), "800")  # 화자 이름 끝
        tabs.append(tab1)
        tab2 = OxmlElement("w:tab")
        tab2.set(qn("w:val"), "left")
        tab2.set(qn("w:pos"), "1500")  # 시간 끝
        tabs.append(tab2)
        pPr.append(tabs)

        # hangging indent (줄바꿈 시 텍스트 열에 맞춤)
        para.paragraph_format.left_indent = Twips(1500)
        para.paragraph_format.first_line_indent = Twips(-1500)

        if line.start_time:
            para.add_run(f"{speaker_name}\t{line.start_time}\t{line.text}")
        else:
            para.add_run(f"{speaker_name}\t\t{line.text}")


def format_docx(
    document: Document,
    speakers: dict,
    lines: list[ScriptLine],
) -> io.BytesIO:
    doc = DocxDocument()
    _add_content(doc, document, speakers, lines)
    buffer = io.BytesIO()  # make empty buffer to save the document
    doc.save(buffer)  # save the document to the buffer
    buffer.seek(0)  # set the pointer to the beginning of the buffer
    return buffer


# TODO: has to be updated to single query with relationships later (next development + tuning phase)
async def make_docx(document_id: int, session: AsyncSession) -> tuple[str, io.BytesIO]:
    result = await session.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Document {document_id} not found",
        )

    speakers_result = await session.execute(
        select(Speaker).where(Speaker.document_id == document_id)
    )
    speakers = {s.id: s.name for s in speakers_result.scalars().all()}

    lines_result = await session.execute(
        select(ScriptLine)
        .where(ScriptLine.document_id == document_id)
        .order_by(ScriptLine.order)
    )
    lines = lines_result.scalars().all()

    title = document.title.replace(".txt", "")
    buffer = format_docx(document, speakers, lines)
    return title, buffer


# TODO: has to be updated to single query with relationships later (next development + tuning phase)
async def make_merged_docx(
    category_ids: list[int], session: AsyncSession
) -> io.BytesIO:
    result = await session.execute(
        select(Document)
        .where(Document.category_id.in_(category_ids))
        .order_by(Document.recorded_date.asc(), Document.recorded_time.asc())
    )
    documents = result.scalars().all()

    doc = DocxDocument()
    for i, document in enumerate(documents):
        if i > 0:
            doc.add_paragraph("")

        speakers_result = await session.execute(
            select(Speaker).where(Speaker.document_id == document.id)
        )
        speakers = {s.id: s.name for s in speakers_result.scalars().all()}

        lines_result = await session.execute(
            select(ScriptLine)
            .where(ScriptLine.document_id == document.id)
            .order_by(ScriptLine.order)
        )
        lines = lines_result.scalars().all()
        _add_content(doc, document, speakers, lines)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
